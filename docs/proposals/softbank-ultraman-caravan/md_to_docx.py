# -*- coding: utf-8 -*-
"""社内共有用MarkdownをWord(.docx)に変換する簡易コンバータ。
usage: python3 md_to_docx.py <input.md> <output.docx>
対応: #/##/### 見出し、- 箇条書き、**太字**、通常段落
"""
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

FONT = "Meiryo"


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)


def add_runs(p, text, size=10.5, base_bold=False):
    # **bold** を分解してランを追加
    for part in re.split(r'(\*\*.+?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_run_font(r, size, bold=True)
        else:
            r = p.add_run(part)
            set_run_font(r, size, bold=base_bold)


def convert(src_path, dst_path):
    doc = Document()
    for line in open(src_path, encoding='utf-8').read().splitlines():
        s = line.rstrip()
        if not s.strip():
            continue
        if s.startswith('# '):
            p = doc.add_paragraph()
            add_runs(p, s[2:].replace('**', ''), size=16, base_bold=True)
            p.paragraph_format.space_after = Pt(10)
        elif s.startswith('## '):
            p = doc.add_paragraph()
            add_runs(p, s[3:].replace('**', ''), size=13, base_bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x23, 0x22, 0x78)
        elif s.startswith('### '):
            p = doc.add_paragraph()
            add_runs(p, s[4:].replace('**', ''), size=11.5, base_bold=True)
        elif s.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, s.strip()[2:])
        elif s.startswith('---'):
            continue
        else:
            p = doc.add_paragraph()
            add_runs(p, s)
            p.paragraph_format.space_after = Pt(4)
    doc.save(dst_path)
    print(f'saved: {dst_path}')


if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
